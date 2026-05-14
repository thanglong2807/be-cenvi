import json
import re
from pathlib import Path
from io import BytesIO
from html.parser import HTMLParser
from sqlalchemy.orm import Session
from fastapi import HTTPException

try:
    from docxtpl import DocxTemplate
    DOCXTPL_AVAILABLE = True
except ImportError:
    DOCXTPL_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Pt, Mm, Inches, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from app.models.contract_model import Contract

TEMPLATES_DIR = Path("static/templates")
TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)


def docx_template_path(template_id: int) -> Path:
    return TEMPLATES_DIR / f"template_{template_id}.docx"


def has_docx_template(template_id: int) -> bool:
    return docx_template_path(template_id).exists()


# ─── HTML → DOCX fallback converter ──────────────────────────────────────────

DEFAULT_FONT = 'Times New Roman'
DEFAULT_SIZE = 12  # pt


def _parse_style(style_str: str) -> dict:
    result = {}
    if not style_str:
        return result
    for prop in style_str.split(';'):
        if ':' in prop:
            k, _, v = prop.partition(':')
            result[k.strip().lower()] = v.strip()
    return result


def _pt_from_css(val: str) -> float | None:
    """Parse CSS length value → points. Handles pt, px, cm, mm, in."""
    val = val.strip()
    for unit, factor in (('pt', 1), ('px', 0.75), ('cm', 28.35), ('mm', 2.835), ('in', 72)):
        if val.endswith(unit):
            try:
                return float(val[:-len(unit)]) * factor
            except ValueError:
                return None
    try:
        return float(val)  # bare number → assume pt
    except ValueError:
        return None


def _align_from_style(style: dict):
    return {
        'center':  WD_ALIGN_PARAGRAPH.CENTER,
        'right':   WD_ALIGN_PARAGRAPH.RIGHT,
        'left':    WD_ALIGN_PARAGRAPH.LEFT,
        'justify': WD_ALIGN_PARAGRAPH.LEFT,  # avoid ugly word spacing
    }.get(style.get('text-align', ''))


def _is_dark_hex(hex6: str) -> bool:
    r, g, b = int(hex6[0:2], 16), int(hex6[2:4], 16), int(hex6[4:6], 16)
    return r < 240 or g < 240 or b < 240


def _prescan_table_info(html: str) -> list[dict]:
    """
    For each <table>, return dict with:
      - col_widths: list[int] of column pixel widths from <col width="X">
      - has_border: bool — True if any <td>/<th> has a non-white border color
    """
    result: list[dict] = []
    stack: list[dict] = []
    for m in re.finditer(r'<(/?)(\w+)([^>]*)>', html, re.IGNORECASE):
        closing, tag, attrs_str = m.group(1) == '/', m.group(2).lower(), m.group(3)
        if tag == 'table' and not closing:
            stack.append({'col_widths': [], 'has_border': False})
        elif tag == 'col' and not closing and stack:
            wm = re.search(r'width=["\']?(\d+)', attrs_str, re.IGNORECASE)
            stack[-1]['col_widths'].append(int(wm.group(1)) if wm else 0)
        elif tag in ('td', 'th') and not closing and stack and not stack[-1]['has_border']:
            sm = re.search(r'style=["\']([^"\']*)["\']', attrs_str, re.IGNORECASE)
            if sm:
                for cm in re.finditer(r'border[^:]*:\s*[^#;]*#([0-9a-fA-F]{6})', sm.group(1), re.IGNORECASE):
                    if _is_dark_hex(cm.group(1)):
                        stack[-1]['has_border'] = True
                        break
        elif tag == 'table' and closing and stack:
            info = stack.pop()
            result.append({'col_widths': info['col_widths'], 'has_border': info['has_border']})
    return result


def _set_table_col_widths(table, col_widths_px: list[int]):
    """Apply pre-scanned pixel widths to table columns."""
    if not col_widths_px:
        return
    cols = table.columns
    for i, w_px in enumerate(col_widths_px):
        if i < len(cols) and w_px > 0:
            cols[i].width = Pt(w_px * 0.75)  # px → pt


def _set_doc_defaults(doc):
    """Set document-level default font and size to match Vietnamese contract style."""
    style = doc.styles['Normal']
    style.font.name = DEFAULT_FONT
    style.font.size = Pt(DEFAULT_SIZE)
    # Also set the rPrDefault in document settings
    try:
        rpr = doc.styles['Normal'].element.get_or_add_rPr()
        rfonts = OxmlElement('w:rFonts')
        rfonts.set(qn('w:ascii'), DEFAULT_FONT)
        rfonts.set(qn('w:hAnsi'), DEFAULT_FONT)
        rpr.insert(0, rfonts)
    except Exception:
        pass


class _HtmlDocxRenderer(HTMLParser):
    HEADING_SIZES = {1: 18, 2: 16, 3: 14, 4: 13, 5: 12, 6: 11}

    def __init__(self, doc, tables_info: list[dict]):
        """
        tables_info: list of dicts per table from _prescan_table_info().
          Each dict: {'col_widths': [int, ...], 'has_border': bool}
        """
        super().__init__()
        self.doc = doc
        self._para = None
        self._fmt: list[dict] = []
        self._tbl_stack: list[dict] = []
        self._tables_info = list(tables_info)
        self._cur_table = self._cur_cell = self._cur_cell_para = None
        self._cur_row_idx = -1
        self._cur_col_idx = 0
        self._list_stack: list[str] = []
        self._list_nums: list[int] = []
        self._li_depth = 0

    def _ensure_para(self):
        if self._cur_cell is not None:
            if self._cur_cell_para is None:
                self._cur_cell_para = self._cur_cell.paragraphs[0]
            return self._cur_cell_para
        if self._para is None:
            self._para = self.doc.add_paragraph()
        return self._para

    def _new_block_para(self, align=None):
        if self._cur_cell is not None:
            if self._cur_cell_para is None:
                self._cur_cell_para = self._cur_cell.paragraphs[0]
            else:
                self._cur_cell_para = self._cur_cell.add_paragraph()
            if align:
                self._cur_cell_para.alignment = align
            return self._cur_cell_para
        self._para = self.doc.add_paragraph()
        if align:
            self._para.alignment = align
        return self._para

    def _merged_fmt(self) -> dict:
        merged = {}
        for f in self._fmt:
            merged.update(f)
        return merged

    def _add_run(self, text: str):
        if not text:
            return
        para = self._ensure_para()
        run = para.add_run(text)
        fmt = self._merged_fmt()
        if fmt.get('bold'):      run.bold = True
        if fmt.get('italic'):    run.italic = True
        if fmt.get('underline'): run.underline = True
        # Always set font — default to Times New Roman if not specified
        run.font.name = fmt.get('font_name') or DEFAULT_FONT
        run.font.size = Pt(fmt['font_size']) if fmt.get('font_size') else Pt(DEFAULT_SIZE)

    def handle_starttag(self, tag, attrs):
        tag = tag.lower()
        attrs = dict(attrs)
        style = _parse_style(attrs.get('style', ''))
        align = _align_from_style(style)

        if tag == 'table':
            info = self._tables_info.pop(0) if self._tables_info else {}
            col_widths = info.get('col_widths', [])
            has_border = info.get('has_border', False)
            col_count = len(col_widths)
            self._tbl_stack.append({
                'table': self._cur_table, 'row_idx': self._cur_row_idx,
                'col_idx': self._cur_col_idx, 'cell': self._cur_cell,
                'cell_para': self._cur_cell_para,
                'col_widths': col_widths,
            })
            if col_count > 0:
                tbl = self.doc.add_table(rows=0, cols=col_count)
                if has_border:
                    try: tbl.style = 'Table Grid'
                    except Exception: pass
                self._cur_table = tbl
                self._pending_col_widths = col_widths
            else:
                self._cur_table = None
                self._pending_col_widths = []
            self._cur_row_idx = -1
            self._cur_col_idx = 0
            self._cur_cell = self._cur_cell_para = None
            return

        if tag in ('colgroup', 'col', 'tbody'):
            return

        if tag == 'tr':
            if self._cur_table is None:
                tbl = self.doc.add_table(rows=0, cols=1)
                self._cur_table = tbl
                self._pending_col_widths = []
            row = self._cur_table.add_row()
            self._cur_row_idx = len(self._cur_table.rows) - 1
            # Apply column widths on first row only
            if self._cur_row_idx == 0 and getattr(self, '_pending_col_widths', []):
                _set_table_col_widths(self._cur_table, self._pending_col_widths)
                self._pending_col_widths = []
            self._cur_col_idx = 0
            self._cur_cell = self._cur_cell_para = None
            return

        if tag in ('td', 'th'):
            if self._cur_table is not None and self._cur_row_idx >= 0:
                row = self._cur_table.rows[self._cur_row_idx]
                if self._cur_col_idx < len(row.cells):
                    self._cur_cell = row.cells[self._cur_col_idx]
                    self._cur_cell_para = None
                    self._cur_col_idx += 1
            return

        if tag in ('p', 'div'):
            if self._li_depth > 0:
                self._fmt.append({})
            else:
                para = self._new_block_para(align)
                # Paragraph indent from margin-left / text-indent
                ml = _pt_from_css(style.get('margin-left', ''))
                ti = _pt_from_css(style.get('text-indent', ''))
                if ml and ml > 1:
                    para.paragraph_format.left_indent = Pt(ml)
                if ti and abs(ti) > 1:
                    para.paragraph_format.first_line_indent = Pt(ti)
                # Line spacing from line-height (numeric ratio)
                lh = style.get('line-height', '')
                try:
                    lh_val = float(lh)
                    if 0.5 < lh_val < 5:
                        from docx.shared import Length
                        from docx.enum.text import WD_LINE_SPACING
                        para.paragraph_format.line_spacing = lh_val
                        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                except (ValueError, ImportError):
                    pass
                self._fmt.append({})
            return

        if tag in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
            level = int(tag[1])
            if self._cur_cell is not None:
                self._cur_cell_para = (self._cur_cell.paragraphs[0]
                                       if self._cur_cell_para is None
                                       else self._cur_cell.add_paragraph())
                if align: self._cur_cell_para.alignment = align
                self._fmt.append({'bold': True, 'font_size': self.HEADING_SIZES.get(level, 12)})
            else:
                self._para = self.doc.add_heading('', level=min(level, 9))
                if align: self._para.alignment = align
                self._fmt.append({'bold': True})
            return

        if tag == 'li':
            self._li_depth += 1
            self._new_block_para(align)
            prefix = ''
            if self._list_stack:
                if self._list_stack[-1] == 'ul':
                    prefix = '• '
                else:
                    n = self._list_nums[-1]
                    self._list_nums[-1] = n + 1
                    prefix = f'{n}. '
            if prefix:
                self._add_run(prefix)
            self._fmt.append({})
            return

        if tag == 'ul':
            self._list_stack.append('ul'); self._list_nums.append(1); return
        if tag == 'ol':
            start = int(attrs.get('start', 1))
            self._list_stack.append('ol'); self._list_nums.append(start); return

        if tag == 'br':
            if self._li_depth == 0:
                self._ensure_para().add_run('\n')
            return

        if tag == 'span':
            fmt: dict = {}
            bw = style.get('font-weight', '')
            if bw in ('700', '800', '900', 'bold', 'bolder'): fmt['bold'] = True
            if style.get('font-style', '') == 'italic':        fmt['italic'] = True
            if 'underline' in style.get('text-decoration', ''): fmt['underline'] = True
            sz = style.get('font-size', '')
            m = re.match(r'(\d+(?:\.\d+)?)\s*pt', sz)
            if m: fmt['font_size'] = float(m.group(1))
            ff = style.get('font-family', '')
            if ff:
                name = ff.split(',')[0].strip().strip('"\'')
                if name: fmt['font_name'] = name
            self._fmt.append(fmt); return

        if tag in ('strong', 'b'): self._fmt.append({'bold': True}); return
        if tag in ('em', 'i'):     self._fmt.append({'italic': True}); return
        if tag == 'u':             self._fmt.append({'underline': True}); return

    def handle_endtag(self, tag):
        tag = tag.lower()
        pop_tags = {'p','div','h1','h2','h3','h4','h5','h6','li','span','strong','b','em','i','u'}
        if tag in pop_tags and self._fmt:
            self._fmt.pop()
        if tag == 'li':
            self._li_depth = max(0, self._li_depth - 1)
        if tag in ('ul', 'ol'):
            if self._list_stack: self._list_stack.pop()
            if self._list_nums:  self._list_nums.pop()
        if tag in ('td', 'th'):
            self._cur_cell = self._cur_cell_para = None
        if tag == 'table' and self._tbl_stack:
            ctx = self._tbl_stack.pop()
            self._cur_table    = ctx['table']
            self._cur_row_idx  = ctx['row_idx']
            self._cur_col_idx  = ctx['col_idx']
            self._cur_cell     = ctx['cell']
            self._cur_cell_para = ctx['cell_para']

    def handle_data(self, data):
        text = re.sub(r'[ \t\r\n]+', ' ', data)
        if not text or text == ' ':
            para = self._ensure_para()
            if text == ' ' and para.runs:
                self._add_run(' ')
            return
        self._add_run(text)

    def handle_entityref(self, name):
        entities = {'nbsp': ' ', 'amp': '&', 'lt': '<', 'gt': '>', 'quot': '"'}
        self._add_run(entities.get(name, ''))

    def handle_charref(self, name):
        try:
            self._add_run(chr(int(name[1:], 16) if name.startswith('x') else int(name)))
        except Exception:
            pass


def _html_to_docx(html: str, doc):
    _set_doc_defaults(doc)
    tables_info = _prescan_table_info(html)
    _HtmlDocxRenderer(doc, tables_info).feed(html)


def _generate_from_html(contract: Contract) -> BytesIO:
    """Fallback: convert HTML template snapshot → DOCX."""
    template_snapshot = json.loads(contract.template_snapshot)
    field_values = json.loads(contract.field_values)
    content = template_snapshot.get('content', '')
    for key, value in field_values.items():
        content = content.replace(f'{{{{{key}}}}}', str(value))
    content = re.sub(r'\{\{[\w_]+\}\}', '', content)

    doc = Document()
    for section in doc.sections:
        section.page_width  = Mm(210)
        section.page_height = Mm(297)
        section.left_margin = section.right_margin = Mm(25)
        section.top_margin  = section.bottom_margin = Mm(20)

    _html_to_docx(content, doc)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ─── Service ──────────────────────────────────────────────────────────────────

class ContractWordService:
    """
    Generates Word (.docx) from contracts.
    - Has .docx template → docxtpl fills {{placeholders}} (preserves exact formatting).
    - No .docx template  → converts HTML template content to DOCX (fallback).
    """

    def generate_word_bytes(self, contract_id: int, db: Session) -> BytesIO:
        contract = db.query(Contract).filter(Contract.id == contract_id).first()
        if not contract:
            raise HTTPException(status_code=404, detail="Contract not found")

        tpl_path = docx_template_path(contract.template_id)

        if tpl_path.exists() and DOCXTPL_AVAILABLE:
            return self._generate_from_docxtpl(contract, tpl_path)
        elif DOCX_AVAILABLE:
            return _generate_from_html(contract)
        else:
            raise HTTPException(status_code=500, detail="python-docx not installed")

    @staticmethod
    def _generate_from_docxtpl(contract: Contract, tpl_path: Path) -> BytesIO:
        field_values = json.loads(contract.field_values)
        tpl = DocxTemplate(str(tpl_path))
        tpl.render(field_values, autoescape=False)
        buf = BytesIO()
        tpl.save(buf)
        buf.seek(0)
        return buf

    @staticmethod
    def save_template_file(template_id: int, content: bytes):
        docx_template_path(template_id).write_bytes(content)

    @staticmethod
    def get_template_file(template_id: int) -> bytes:
        path = docx_template_path(template_id)
        if not path.exists():
            raise HTTPException(status_code=404, detail="Word template chưa được upload")
        return path.read_bytes()

    @staticmethod
    def delete_template_file(template_id: int):
        path = docx_template_path(template_id)
        if path.exists():
            path.unlink()
